"""
One-off script to generate docs/KRIME_AI_Demo_Guide.docx from DEMO_GUIDE.md,
embedding the real screenshots from docs/demo-screenshots/.
Not part of the app runtime -- run manually whenever DEMO_GUIDE.md changes.
"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(__file__)
SHOTS = os.path.join(HERE, "demo-screenshots")
OUT = os.path.join(HERE, "KRIME_AI_Demo_Guide.docx")

NAVY = RGBColor(0x0D, 0x1B, 0x2A)
ACCENT = RGBColor(0x1A, 0x73, 0xE8)
GREY = RGBColor(0x55, 0x55, 0x55)

doc = Document()

# ── Base style tweaks ──────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)

def add_title_page():
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🚔 KRIME AI")
    run.font.size = Pt(40)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Hackathon Demo Guide")
    r2.font.size = Pt(22)
    r2.font.color.rgb = NAVY

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Datathon 2026 · Challenge 1 · Zoho Catalyst")
    r3.font.size = Pt(13)
    r3.italic = True
    r3.font.color.rgb = GREY

    doc.add_paragraph()
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run(
        "A step-by-step script for demoing KRIME AI to judges,\n"
        "with real screenshots of every feature."
    )
    r4.font.size = Pt(11)
    r4.font.color.rgb = GREY

    doc.add_paragraph()
    links = [
        ("Live App", "https://krime-ai-60078097690.development.catalystserverless.in/app/index.html"),
        ("Slate Frontend", "https://krime-ai-slate-jjmjpohd.onslate.in/"),
        ("Catalyst Console", "https://console.catalyst.zoho.in/baas/60078097690/index#/"),
    ]
    for label, url in links:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f"{label}:  ")
        r.bold = True
        r2 = p.add_run(url)
        r2.font.color.rgb = ACCENT
    doc.add_page_break()

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h

def add_body(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    return p

def add_bullets(items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)

def add_numbered(items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)

def add_image(filename, caption=None, width_in=6.0):
    path = os.path.join(SHOTS, filename)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width_in))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        r.italic = True
        r.font.size = Pt(9)
        r.font.color.rgb = GREY

def add_table(headers, rows, header_fill="1A73E8"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(hdr_cells[i], header_fill)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table

def add_divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("─" * 60)
    r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)


# ══════════════════════════ TITLE PAGE ══════════════════════════
add_title_page()

# ══════════════════════════ DEMO FLOW TABLE ══════════════════════════
add_heading("⏱ Suggested Demo Flow (8–10 minutes)", level=1)
add_table(
    ["#", "Feature", "Time"],
    [
        ["1", "Secure Sign-In / Role-Based Access", "1 min"],
        ["2", "Conversational AI Chat (NL→SQL + Context)", "2 min"],
        ["3", "Dashboard & KPIs", "1 min"],
        ["4", "Predictive Analytics + AI Anomaly Detection", "1.5 min"],
        ["5", "Crime Heatmap", "0.5 min"],
        ["6", "Criminal Network Visualization", "1 min"],
        ["7", "Document Intelligence (RAG)", "1 min"],
        ["8", "Audit Trail / Explainability", "0.5 min"],
        ["9", "Voice Interaction + Role Restriction Demo", "1 min"],
    ],
)

# ══════════════════════════ SECTION 1 ══════════════════════════
add_heading("1️⃣ Secure Sign-In · Role-Based Access", level=1)
add_body(
    "Open the live app URL. Judges will see a gated login screen — "
    "nothing is accessible without authentication."
)
add_image("01-login.png", "Login screen")
add_body("Talking points:", bold=True)
add_bullets([
    "4 distinct roles: Admin, SP, Inspector, Analyst — each with a different permission slice.",
    "Demo credentials are shown directly on the screen for judges to try themselves.",
    "Enforcement happens server-side (require_role() decorator on every Flask route) — "
    "not just hidden UI, so it can't be bypassed by calling the API directly.",
])
add_body("Try it: Sign in as admin / Admin@123 for full access.", italic=True)
add_divider()

# ══════════════════════════ SECTION 2 ══════════════════════════
add_heading("2️⃣ Conversational AI Chat — NL → SQL + Context Awareness", level=1)
add_body("After signing in, you land on the AI Chat tab with example queries.")
add_image("02-chat-welcome.png", "Chat welcome screen")
add_body('Type or click a suggestion, e.g. "Show crimes by district". The AI:')
add_numbered([
    "Detects intent → builds parameterized SQL",
    "Executes against the SQLite database (1,500 synthetic FIR records)",
    "Generates a natural-language insight via Catalyst QuickML (GLM-4.7-Flash)",
    "Shows the exact SQL query for transparency",
])
add_image("03-chat-response-chart.png", "Chat response with chart, table, and AI insight")
add_body("Talking points:", bold=True)
add_bullets([
    "Notice the chart panel auto-renders alongside the answer.",
    "The AI Insight line is LLM-generated commentary on the data — not hardcoded.",
    "The chat stays clean and end-user friendly by design — the exact SQL executed for every "
    "answer is still fully available (untruncated) in the Audit Trail tab for full "
    "explainability, without cluttering the conversation.",
])
add_body(
    'Context-awareness demo: Ask a follow-up like "What about theft?" without repeating '
    "the district — the AI remembers what you were just discussing and carries the context "
    'forward (visible later in the Audit Trail tab with a "Context carried over" badge).',
    italic=True,
)
add_body(
    "Bilingual demo: Switch to Kannada and ask "
    "\u0cac\u0cc6\u0c82\u0c97\u0cb3\u0ccd\u0cc2\u0cb0\u0cbf\u0aa8\u0cb2\u0ccd\u0cb2\u0cbf "
    "\u0ca8\u0cc6\u0c95\u0ccd\u0cb7\u0cb0\u0cbf\u0cb8\u0cb2\u0ccb\u0c95 \u0c05\u0cae\u0cb0\u0cbf\u0c95 "
    "\u0c8e\u0cb7\u0ccd\u0cb2\u0cbf \u0ca4\u0cbf\u0cb3\u0cc6\u0cb8\u0cbf",
    italic=True,
)
add_divider()

# ══════════════════════════ SECTION 3 ══════════════════════════
add_heading("3️⃣ Dashboard — Key Performance Indicators", level=1)
add_body("Click Dashboard in the sidebar.")
add_image("04-dashboard.png", "Dashboard with KPIs and charts")
add_body("Talking points:", bold=True)
add_bullets([
    "8 KPI cards: total cases, pending investigations, wanted accused, clearance rate, "
    "total accused/victims, station count, top crime type.",
    "6 interactive Chart.js visualizations: yearly trend, crime type distribution, case status, "
    "top districts, arrest status, severity distribution.",
    "All data-driven from the live SQLite database — click Reload DB (Admin only) to reseed "
    "with fresh synthetic data live during the demo if desired.",
])
add_divider()

# ══════════════════════════ SECTION 4 ══════════════════════════
add_heading("4️⃣ Predictive Analytics + AI Anomaly Detection", level=1)
add_body("Click Analytics.")
add_image("05-analytics.png", "Analytics tab with predictive insights")
add_body("Talking points:", bold=True)
add_bullets([
    "AI Predictive Insights box: QuickML-generated commentary on year-over-year trends, "
    "flagging upward/downward trajectories as early warnings.",
    'Scroll down and click "Scan Stations" to trigger the custom QuickML AutoML pipeline — '
    "a trained classification model that flags police stations whose case volume/severity "
    "statistically deviates from normal.",
])
add_image("06-anomaly-detection.png", "Anomaly detection results flagging 2 stations")
add_body("Talking points:", bold=True)
add_bullets([
    'This directly satisfies the "custom pipeline builders for anomaly or fraud detection '
    'models" requirement — it\'s a real trained AutoML endpoint, not a mock.',
    "Flagged stations (e.g. Bajpe, Ullal in Mangaluru) show case count, average severity, "
    "pending count, and high-severity count — actionable for administrative review.",
])
add_divider()

# ══════════════════════════ SECTION 5 ══════════════════════════
add_heading("5️⃣ Crime Heatmap", level=1)
add_body("Click Heatmap.")
add_image("07-heatmap.png", "Crime heatmap of Karnataka")
add_body("Talking points:", bold=True)
add_bullets([
    "Custom canvas-rendered heatmap of Karnataka with major cities labeled.",
    "Filter by crime type and year using the dropdowns — the glow intensity/color encodes severity.",
    "500 incident points rendered live from /api/heatmap.",
])
add_divider()

# ══════════════════════════ SECTION 6 ══════════════════════════
add_heading("6️⃣ Criminal Network Visualization", level=1)
add_body("Click Network (visible only to Admin/SP/Inspector — hidden for Analyst).")
add_image("08-network.png", "Criminal network graph")
add_body("Talking points:", bold=True)
add_bullets([
    "Interactive force-directed graph (vis-network) showing co-accused sharing gang affiliations.",
    "Red nodes = Wanted/Absconding, Blue nodes = Arrested.",
    "Click any node to see full criminal profile details and connection count.",
    "This is the feature explicitly role-gated — demonstrates how sensitive identity-linked "
    "data is restricted from the Analyst role.",
])
add_divider()

# ══════════════════════════ SECTION 7 ══════════════════════════
add_heading("7️⃣ Document Intelligence — RAG over the KSP SOP Manual", level=1)
add_body("Click Documents.")
add_image("09-documents-rag.png", "RAG document Q&A with cited sources")
add_body("Talking points:", bold=True)
add_bullets([
    "Ask natural-language questions about the KSP SOP Investigation Manual "
    "(uploaded to Catalyst's QuickML Knowledge Base).",
    "Answers are grounded — the response cites the exact source document and snippet, "
    "not a hallucinated answer.",
    'Try: "What is the case clearance rate target?" → returns the real policy figure with a citation.',
])
add_divider()

# ══════════════════════════ SECTION 8 ══════════════════════════
add_heading("8️⃣ Audit Trail — Explainable AI", level=1)
add_body("Click Audit Trail.")
add_image("10-audit-trail.png", "Audit trail showing SQL and context tracking")
add_body("Talking points:", bold=True)
add_bullets([
    "Every single query is logged with: original question, detected intent, exact SQL executed, "
    "and result count.",
    'This is the "explainable AI with audit trails" requirement — nothing is a black box; '
    "every answer can be traced back to its exact query.",
    "Context-carried-over queries show a badge indicating which entity (district/year/crime type) "
    "was remembered from earlier in the conversation.",
])
add_divider()

# ══════════════════════════ SECTION 9 ══════════════════════════
add_heading("9️⃣ Voice Interaction + Role-Based Restriction (side-by-side)", level=1)
add_body(
    "Voice input/output: Click the mic button in the chat toolbar and speak a query "
    "(English or Kannada supported via en-IN/kn-IN). It auto-sends once you stop talking. "
    "The AI's response is read back aloud automatically — toggle with the speaker button."
)
add_body("Role restriction: Log out and sign back in as analyst1 / Analyst@123. Notice:")
add_bullets([
    "The Network tab has disappeared from the sidebar.",
    "The Export PDF button is gone from the chat toolbar and sidebar footer.",
    "The Reload DB button is gone (Admin-only).",
])
add_image("11-analyst-role-restricted.png", "Analyst role with restricted UI — no Network tab, no Export PDF button")
add_body("Talking points:", bold=True)
add_bullets([
    "Compare the sidebar icon count: 7 icons for Admin/SP/Inspector vs 6 icons for Analyst "
    "(Network hidden).",
    "This isn't just cosmetic — calling /api/network directly as an Analyst returns an HTTP "
    "403 Forbidden from the backend, proving the restriction is enforced server-side, "
    "not just hidden in the UI.",
])
add_divider()

# ══════════════════════════ ARCHITECTURE REFERENCE ══════════════════════════
add_heading("🏗 Architecture Reference", level=1)
add_body(
    "See README.md (Architecture section) for the full system architecture diagram covering "
    "the frontend (Catalyst Slate + Web Client), backend (Flask Advanced I/O function), "
    "data layer (SQLite), and the three QuickML AI integrations (GLM chat insights, "
    "RAG document Q&A, AutoML anomaly detection)."
)
arch_path = os.path.join(HERE, "images", "architecture.svg")
add_body(
    "(Architecture diagram is provided separately as docs/images/architecture.svg "
    "in the project repository — SVG format is not embedded in this Word document.)",
    italic=True,
)
add_divider()

# ══════════════════════════ QUICK REFERENCE ══════════════════════════
add_heading("🔑 Quick Reference — Demo Accounts", level=1)
add_table(
    ["Role", "Username", "Password", "Notable Restrictions"],
    [
        ["Admin", "admin", "Admin@123", "None — full access incl. DB reseed"],
        ["SP", "sp.blru", "Sp@12345", "None — full investigative access"],
        ["Inspector", "inspector1", "Inspector@123", "No DB reseed"],
        ["Analyst", "analyst1", "Analyst@123", "No Network graph, PDF export, Anomaly scan, DB reseed"],
    ],
)

add_heading("💬 Sample Queries to Showcase Range", level=1)
add_table(
    ["Category", "Query"],
    [
        ["Aggregate", "How many total crimes are registered in Karnataka?"],
        ["District breakdown", "Show crimes by district"],
        ["Trend", "Show crime trend from 2019 to 2025"],
        ["Hotspot", "Top crime hotspot stations"],
        ["Context follow-up", "How many crimes in Mysuru? → then What about theft?"],
        ["Predictive", "Predict crime trend for next months"],
        ["Kannada", "\u0cac\u0cc6\u0c82\u0c97\u0cb3\u0ccd\u0cc2\u0cb0\u0cbf\u0aa8\u0cb2\u0ccd\u0cb2\u0cbf \u0ca8\u0cc6\u0c95\u0ccd\u0cb7\u0cb0\u0cbf\u0cb8\u0cb2\u0ccb\u0c95 \u0c05\u0cae\u0cb0\u0cbf\u0c95 \u0c8e\u0cb7\u0ccd\u0cb2\u0cbf \u0ca4\u0cbf\u0cb3\u0cc6\u0cb8\u0cbf"],
        ["RAG", "What is the case clearance rate target? (Documents tab)"],
    ],
)

doc.save(OUT)
print(f"Saved: {OUT}")
